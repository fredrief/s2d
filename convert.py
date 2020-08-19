import os
import shutil

from text_utils import *
from docx import Document
from docx.shared import Cm

from image_utils import *


def convert(pdf_doc_filename, docx_filename):
    doc = fitz.open(pdf_doc_filename) # PDF document
    word_doc = Document() # Docx document

    # Set paragraph settings
    paragraph_format = word_doc.styles['Normal'].paragraph_format
    paragraph_format.space_after = 0

    ############ CONVERT MAIN PAGE ##############
    main_page = doc[0]
    txtblocks, imblocks = get_blocks(main_page)
    # Insert Heading
    heading = get_page_heading(txtblocks, 0)
    if heading:
        word_doc.add_heading(heading, 0)
    # Insert Images
    im_par = word_doc.add_paragraph()
    im_run = im_par.add_run()
    main_page_images = get_images(main_page, imblocks, 0)
    for image in main_page_images:
        if image[1] == "H":
            im_run.add_picture(image[0], height=Cm(5))
        else:
            im_run.add_picture(image[0], width=Cm(5))
        im_run.add_text(" ")
    # Insert Text
    paragraphs = get_paragraphs(txtblocks)
    for p in paragraphs:
        word_doc.add_paragraph(p[0])

    ############ CONVERT OTHER PAGES ##################
    for page_index in range(len(doc)):
        if 0<page_index:
            print("{}%".format(math.ceil(page_index/len(doc)*100)))
            page = doc[page_index]
            txtblocks, imblocks = get_blocks(page)

            # Insert Heading
            heading = get_page_heading(txtblocks, page_index)
            if heading:
                word_doc.add_heading(heading, 2)

            # Insert Images
            im_par = word_doc.add_paragraph()
            im_run = im_par.add_run()
            page_images = get_images(page, imblocks, page_index)
            for image in page_images:
                if image[1] == "H":
                    im_run.add_picture(image[0], height=Cm(5))
                else:
                    im_run.add_picture(image[0], width=Cm(5))
                im_run.add_text(" ")

            # Insert Text
            paragraphs = get_paragraphs(txtblocks)
            for p in paragraphs:
                # Check which level the bullet point is
                if p[1] == 0:
                    word_doc.add_paragraph(p[0], style="List Bullet")
                elif p[1] == 1:
                    # Level 1
                    word_doc.add_paragraph(p[0], style="List Bullet 2")
                elif p[1] == 2:
                    # Level 2
                    word_doc.add_paragraph(p[0], style="List Bullet 3")

    word_doc.save(docx_filename)

def get_blocks(page):
    blocks = page.getText("dict")["blocks"]
    imblocks = []
    textblocks = []
    for block in blocks:
        if block["type"] == 0:
        # then the block has text
            textblocks.append(block)
        else:
            imblocks.append(block)
    return textblocks, imblocks

def delete_images():
    folder = 'images'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))